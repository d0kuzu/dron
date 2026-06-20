import os
import subprocess
import json
import copy
import re
import io
import docx
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from dotenv import load_dotenv

load_dotenv()

# Import config helpers from main
from main import CONFIG_PATH, DEFAULT_CONFIG, load_config, _deep_merge

app = FastAPI()

@app.get("/")
async def root():
    # Serve index.html
    return FileResponse("static/index.html")

@app.get("/drone-data")
async def get_data():
    matched = {"opportunities": []}
    if os.path.exists("sam_matches.json"):
        try:
            with open("sam_matches.json", "r", encoding="utf-8") as f:
                matched = json.load(f)
        except Exception:
            pass

    excluded = {"opportunities": []}
    if os.path.exists("sam_excluded.json"):
        try:
            with open("sam_excluded.json", "r", encoding="utf-8") as f:
                excluded = json.load(f)
        except Exception:
            pass
            
    return {
        "matched": matched,
        "excluded": excluded
    }

@app.get("/drone-config")
async def get_config():
    """Return the current drone configuration (merged with defaults)."""
    cfg = load_config()
    return {"config": cfg, "defaults": copy.deepcopy(DEFAULT_CONFIG)}

@app.post("/drone-config")
async def save_config(payload: dict):
    """Save updated drone configuration to drone_config.json."""
    config_data = payload.get("config")
    if config_data is None:
        raise HTTPException(status_code=400, detail="Missing 'config' key in request body.")
    
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(config_data, f, ensure_ascii=False, indent=2)
        
        # Reload module-level config in main.py
        import main
        main._CFG = load_config()
        main.SAM_SEARCH_URL_DEFAULTS = tuple(main._CFG["search"]["sam_search_urls"])
        main.DEFAULT_PTYPES = tuple(main._CFG["search"]["procurement_types"])
        main.TARGET_DEPARTMENT_KEYWORDS = main._CFG["keywords"]["target_departments"]
        main.RFP_TOPIC_KEYWORDS = main._CFG["keywords"]["rfp_topics"]
        main.KNOWN_COMPETITOR_BRANDS = main._CFG["competitors"]["brands"]
        main.OR_EQUAL_PATTERNS = main._CFG["competitors"]["or_equal_patterns"]
        main.NEGATIVE_KEYWORDS = main._CFG["keywords"]["negative"]
        main.COMPONENTS_ONLY_KEYWORDS = main._CFG["keywords"]["components_only"]
        main.PAYLOAD_PATTERNS = main._CFG["patterns"]["payload"]
        main.WINCH_PATTERNS = main._CFG["patterns"]["winch"]
        main.PARACHUTE_PATTERNS = main._CFG["patterns"]["parachute"]
        main.COMMS_PATTERNS = main._CFG["patterns"]["comms"]
        main.DEFAULT_NAICS_CODES = main._CFG["search"]["naics_codes"]
        main.SMALL_BIZ_CODES = set(main._CFG["small_biz_codes"])
        
        return {"status": "ok", "config": load_config()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/parse-docx")
async def parse_docx(file: UploadFile = File(...)):
    """Parse an uploaded DOCX file for drone characteristics."""
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Только файлы .docx поддерживаются.")
    
    try:
        content = await file.read()
        doc = docx.Document(io.BytesIO(content))
        
        extracted = {}
        found_weights = []
        found_comms = []
        has_winch = False
        has_parachute = False
        
        # Parse all tables in the document
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    key = row.cells[0].text.strip().lower()
                    val = row.cells[1].text.strip().lower()
                    
                    if "payload" in key and "max" in key:
                        match = re.search(r'(\d+(?:\.\d+)?)\s*kg', val)
                        if match: found_weights.append(float(match.group(1)))
                    elif "delivery weight" in key:
                        match = re.search(r'(\d+(?:\.\d+)?)\s*lbs', val)
                        if match: found_weights.append(float(match.group(1)) * 0.453)
                    elif "connectivity" in key:
                        if "lte" in val: found_comms.append("lte")
                        if "900mhz" in val or "900 mhz" in val: found_comms.append("900mhz")
                    elif "payload release" in key or "winch" in key:
                        has_winch = True
                    elif "parachute" in key or "recovery system" in key:
                        has_parachute = True
                    
        # Extract keywords directly from the text as fallback
        for para in doc.paragraphs:
            text = para.text.lower()
            if "winch" in text or "gravity release" in text:
                has_winch = True
            if "parachute" in text or "recovery system" in text:
                has_parachute = True
            if "lte" in text: found_comms.append("lte")
            if "900mhz" in text or "900 mhz" in text: found_comms.append("900mhz")
        
        if found_weights:
            extracted["max_payload_kg"] = round(max(found_weights), 1)
        
        if found_comms:
            extracted["comms"] = list(set(found_comms))
            
        if has_winch:
            extracted["has_winch"] = True
            
        if has_parachute:
            extracted["has_parachute"] = True
            
        return {"status": "ok", "extracted": extracted}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка парсинга файла: {str(e)}")

@app.post("/drone-refresh")
async def refresh_data():
    api_key = os.getenv("SAM_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="SAM_API_KEY is not configured in .env file.")
        
    try:
        # Run the existing main.py script
        result = subprocess.run(
            ["python", "main.py", "--api-key", api_key],
            capture_output=True,
            text=True
        )
        if result.returncode != 0:
            raise Exception(f"Script failed: {result.stderr}")
            
        # If successful, read the newly generated data
        return await get_data()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/test-parse-local-docx")
async def test_parse_local_docx():
    """Test endpoint that reads the local characteristics.docx and returns extracted JSON data."""
    file_path = "characteristics.docx"
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Файл characteristics.docx не найден в корне проекта.")
        
    try:
        doc = docx.Document(file_path)
        
        extracted = {}
        found_weights = []
        found_comms = []
        has_winch = False
        has_parachute = False
        
        # Parse all tables in the document
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    key = row.cells[0].text.strip().lower()
                    val = row.cells[1].text.strip().lower()
                    
                    if "payload" in key and "max" in key:
                        match = re.search(r'(\d+(?:\.\d+)?)\s*kg', val)
                        if match: found_weights.append(float(match.group(1)))
                    elif "delivery weight" in key:
                        match = re.search(r'(\d+(?:\.\d+)?)\s*lbs', val)
                        if match: found_weights.append(float(match.group(1)) * 0.453)
                    elif "connectivity" in key:
                        if "lte" in val: found_comms.append("lte")
                        if "900mhz" in val or "900 mhz" in val: found_comms.append("900mhz")
                    elif "payload release" in key or "winch" in key:
                        has_winch = True
                    elif "parachute" in key or "recovery system" in key:
                        has_parachute = True
                    
        # Extract keywords directly from the text as fallback
        for para in doc.paragraphs:
            text = para.text.lower()
            if "winch" in text or "gravity release" in text:
                has_winch = True
            if "parachute" in text or "recovery system" in text:
                has_parachute = True
            if "lte" in text: found_comms.append("lte")
            if "900mhz" in text or "900 mhz" in text: found_comms.append("900mhz")
        
        if found_weights:
            extracted["max_payload_kg"] = round(max(found_weights), 1)
        
        if found_comms:
            extracted["comms"] = list(set(found_comms))
            
        if has_winch:
            extracted["has_winch"] = True
            
        if has_parachute:
            extracted["has_parachute"] = True
            
        return {
            "status": "ok", 
            "message": "Это тестовый вывод (ничего не сохранено)",
            "extracted": extracted
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка парсинга файла: {str(e)}")
